"""Write the sample-size / regression recap to a Word document — FINAL clean-sample version."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from getpass import getuser

user = getuser()
OUT = f'C:/Users/{user}/Documents/GitHub/tennis-homophily/recap_lingqing_20260518.docx'

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'BDD7EE')
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

# ════════════════════════════════════════════════════════════════════════════════
# Title
# ════════════════════════════════════════════════════════════════════════════════
title = doc.add_heading('Tennis Homophily — Final Clean-Sample Recap', 0)
title.runs[0].font.size = Pt(16)
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
doc.add_paragraph().add_run(
    'Updated addressing Lingqing\'s review comments — 18 May 2026 (final clean sample)'
).font.italic = True
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════════
h1('0. Data Quality Fixes — What Changed and Why')
# ════════════════════════════════════════════════════════════════════════════════

body(
    'After Lingqing\'s initial review we identified three coding errors in the '
    'nationality / language variables. All three are now fixed; the regression '
    'sample shrinks from N ≈ 3,446 to N = 3,138 team-obs.'
)
doc.add_paragraph()

h2('0a. Bug 1 — NaN nationality coded as "different"')
body(
    '38 players had no scraped country (Wikipedia page missing or disambiguation '
    'failure). The original code evaluated (NaN == "Spain") → False → 0, '
    'treating a missing country as proof that the two players are different '
    'nationalities. Fix: same_country = NaN whenever either player\'s country '
    'is unknown; those matches are now excluded from the regression.'
)

h2('0b. Bug 2 — NaN language coded as "different"')
body(
    'The language variables were populated via a left-join on the CEPII bilateral '
    'gravity dataset (iso3_o, iso3_d). Unmatched rows were filled with fillna(0), '
    'treating "not found in lookup" as "different language." This affected: '
    '(i) country pairs absent from the bilateral file, and '
    '(ii) players whose country code was missing entirely.'
)
body(
    'Fix: (1) same-iso3 pairs → same_language = 1 and ling_prox = 1 by definition '
    '(CEPII excludes self-joins); (2) both iso3 known but pair not in lookup → 0; '
    '(3) either iso3 missing → NaN → match excluded.'
)

h2('0c. Bug 3 — Missing iso3 codes for 22 countries')
body(
    'The player scraper populated country_std but left iso3 blank for 22 nationalities '
    '(e.g., Georgia, Switzerland, Hungary, South Korea, Norway, Lebanon). '
    'Fix: hardcoded ISO3_MAP applied before the language merge, recovering iso3 for '
    'all affected players.'
)

doc.add_paragraph()
add_table(
    ['Issue', 'Matches affected', 'Resolution'],
    [
        ('NaN same_country coded as 0',        '~100 matches', 'np.where guard → NaN → excluded'),
        ('NaN language/prox coded as 0',       '~170 matches', 'Remove fillna(0); keep NaN → excluded'),
        ('Missing iso3 (22 countries)',         '~60 matches',  'ISO3_MAP hardcoded in cleaning pipeline'),
        ('Same-country pairs absent from CEPII', 'all same-nat pairs', 'Explicit same-iso3 → 1 before lookup'),
    ],
    col_widths=[2.4, 1.6, 2.6]
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════════
h1('1. Table 3 — Match Win: final sample funnel')
# ════════════════════════════════════════════════════════════════════════════════

body(
    'The complete sample funnel from raw data to the regression sample is:'
)
doc.add_paragraph()
add_table(
    ['Stage', 'N'],
    [
        ('Raw dataset',                                              '1,840'),
        ('After dropping retirements / walkovers (47)',             '1,793'),
        ('  of which: Grand Slam matches',                          '1,730'),
        ('  of which: Olympics matches (excluded from Table 3)',    '63'),
        ('GS matches dropped: ranking incomplete (both players unranked)', '14'),
        ('GS matches dropped: nationality / language missing for ≥1 player', '129'),
        ('GS regression sample (Table 3)',                          '1,598'),
        ('Team-obs (2 per match, 0 orphans)',                       '3,196'),
    ],
    col_widths=[4.0, 1.8]
)
doc.add_paragraph()

bullet(
    '1,793 − 1,730 = 63 are Olympics matches, not missing data.',
    bold_prefix='Olympics: '
)
bullet(
    '14 GS matches dropped because both players on one team lack a valid '
    'doubles ranking. 129 GS matches dropped because at least one player\'s '
    'country or iso3 is unresolvable after all fixes.',
    bold_prefix='Dropped: '
)
bullet(
    'N = 3,196 team-obs from 1,598 GS matches. '
    'Zero singleton observations (all remaining matches have complete data '
    'for both teams).',
    bold_prefix='Final: '
)

doc.add_paragraph()

h2('1b. Updated Table 3 regression results')
body(
    'With the corrected sample all three culture measures are statistically '
    'indistinguishable from zero (previously same_language was p = 0.023 with '
    'the mislabelled NaN observations inflating the estimate).'
)
doc.add_paragraph()
add_table(
    ['Culture variable', 'Coef', 'SE', 'p-value', 'Sig.'],
    [
        ('Same nationality',       '+0.058', '0.079', '0.466', ''),
        ('Same official language', '+0.116', '0.077', '0.129', ''),
        ('Language proximity',     '+0.097', '0.077', '0.205', ''),
    ],
    col_widths=[2.4, 0.8, 0.8, 0.9, 0.7]
)
doc.add_paragraph()
body(
    'Controls: team avg. doubles rank (−0.0041***), opponent avg. rank (+0.0033***), '
    'teammate rank gap, top-100 singles indicator. FE: tournament × year, round. '
    'Clustered SE by match.',
    italic=True
)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════════
h1('2. Table 4 — Tiebreak Win: updated sample counts')
# ════════════════════════════════════════════════════════════════════════════════

body(
    'The tiebreak sample draws from the same 1,569-match clean GS panel. '
    'All orphan (singleton) observations are eliminated.'
)

h2('2a. Sample counts by tiebreak type')
add_table(
    ['Sample', 'Matches', 'Team-obs (N)', 'Orphans'],
    [
        ('7-pt regular tiebreaks (sets 1–2)',   '712', '1,424', '0'),
        ('Any tiebreak (all types, reference)', '797', '1,594', '0'),
    ],
    col_widths=[2.8, 1.0, 1.4, 0.9]
)
doc.add_paragraph()

h2('2b. Updated regression results (7-pt regular tiebreaks, primary)')
add_table(
    ['Culture variable', 'Coef', 'SE', 'p-value', 'Sig.'],
    [
        ('Same nationality',       '+0.104', '0.120', '0.384', ''),
        ('Same official language', '+0.059', '0.116', '0.610', ''),
        ('Language proximity',     '+0.041', '0.116', '0.723', ''),
    ],
    col_widths=[2.4, 0.8, 0.8, 0.9, 0.7]
)
doc.add_paragraph()
bullet(
    'No homophily advantage in tiebreaks. Including all tiebreak types '
    '(N = 1,594 obs, 797 matches) gives near-identical coefficients '
    '(+0.024, +0.023, +0.001; all p > 0.80).',
    bold_prefix='Finding: '
)
bullet(
    'The previous any-tiebreak sample (N = 1,720) included 6 orphan obs and mixed '
    'GS + Olympics matches; the new 7-pt GS-only sample (N = 1,424) is cleaner '
    'and yields the same null result.',
    bold_prefix='Change from prior version: '
)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════════
h1('3. Table 5 — Comeback Win: adversity sample update')
# ════════════════════════════════════════════════════════════════════════════════

h2('3a. Sample definition (unchanged structure, updated N)')
body(
    'Adversity_im = 1 if team i lost set 1 in match m. '
    'Outcome: win = 1 (comeback) or 0 (any loss after losing set 1). '
    'Each match contributes exactly one observation.'
)

add_table(
    ['Scenario', 'Set 1', 'Set 2', 'Set 3', 'Outcome', 'N', 'Share'],
    [
        ('Comeback win',         'Loss', 'Win',  'Win',  'win = 1', '310', '19.4%'),
        ('3-set loss',           'Loss', 'Win',  'Loss', 'win = 0', '392', '24.5%'),
        ('Straight-set loss',    'Loss', 'Loss', '—',    'win = 0', '896', '56.1%'),
        ('Total adversity obs',  '',     '',     '',     '',        '1,598', '100%'),
    ],
    col_widths=[1.8, 0.7, 0.7, 0.7, 0.8, 0.7, 0.8]
)
doc.add_paragraph()
body(
    'Standard errors: HC3 robust (no within-match clustering — 1 obs per match). '
    'FE: tournament × year + round. Tournament × year cells with no outcome variation dropped.',
    italic=True
)

h2('3b. Updated regression results (N = 1,569)')
add_table(
    ['Culture variable', 'Coef', 'SE', 'p-value', 'Sig.'],
    [
        ('Same nationality',       '−0.235', '0.137', '0.085', '*'),
        ('Same official language', '−0.324', '0.132', '0.014', '**'),
        ('Language proximity',     '−0.330', '0.133', '0.013', '**'),
    ],
    col_widths=[2.4, 0.8, 0.8, 0.9, 0.7]
)
doc.add_paragraph()

h2('3c. Comparison: prior version vs final clean sample')
add_table(
    ['Culture variable', 'Prior version (N = 1,717)', 'Final clean (N = 1,598)'],
    [
        ('Same nationality',       '−0.192 (p = 0.147, n.s.)', '−0.235* (p = 0.085)'),
        ('Same official language', '−0.270** (p = 0.035)',      '−0.324** (p = 0.014)'),
        ('Language proximity',     '−0.263** (p = 0.040)',      '−0.330** (p = 0.013)'),
    ],
    col_widths=[2.2, 2.2, 2.2]
)
doc.add_paragraph()

bullet(
    'The negative sign is preserved: homophily teams that face adversity '
    '(lost set 1) are less likely to stage a comeback.',
    bold_prefix='Key finding: '
)
bullet(
    'Same_language and ling_prox remain significant and are more precise '
    '(p = 0.013/0.014 vs 0.035/0.040) after removing the NaN-coded-as-0 observations '
    'that diluted the effect. Same_nationality gains a * (p = 0.085) in the final clean sample.',
    bold_prefix='Effect of data cleaning: '
)
bullet(
    'The smaller N (1,598 vs 1,717) reflects 119 adversity observations removed '
    'because the team\'s nationality / language data was previously silently '
    'mislabelled as "different" (NaN → 0).',
    bold_prefix='Sample change: '
)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════════
h1('Summary table: N across all three regressions (final clean sample)')
# ════════════════════════════════════════════════════════════════════════════════
add_table(
    ['Table', 'Outcome', 'Sample', 'Matches', 'Team-obs (N)', 'Orphans'],
    [
        ('Table 3', 'Match win',              'GS only, complete matches',                '1,598', '3,196', '0'),
        ('Table 4A', 'Tiebreak win (7-pt)',   'GS, regular_tb = 1 (sets 1–2)',            '712',   '1,424', '0'),
        ('Table 4B', 'Tiebreak win (any)',    'GS, any_tb = 1 (reference)',               '797',   '1,594', '0'),
        ('Table 5', 'Comeback win',           'GS, lost set 1 (adversity, 1 obs/match)', '1,598', '1,598', 'n/a'),
    ],
    col_widths=[0.85, 1.4, 2.5, 0.85, 1.0, 0.75]
)

doc.add_paragraph()
doc.add_paragraph()
body(
    'Notebook updated: code/cleaning/final_ds.ipynb (ISO3_MAP, same_country NaN guard, '
    'language fill logic) and code/analysis/homophily_analysis.ipynb (cells 14 & 16: '
    'match-level drop for incomplete ranking or nationality/language).',
    italic=True
)

doc.save(OUT)
print(f'Saved: {OUT}')
